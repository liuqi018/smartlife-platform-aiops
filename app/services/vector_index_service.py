"""向量索引服务模块"""

import hashlib
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Optional

from loguru import logger

from app.core.milvus_client import milvus_manager
from app.services.document_splitter_service import document_splitter_service
from app.services.vector_store_manager import vector_store_manager


class IndexingResult:
    """索引结果类"""

    def __init__(self):
        self.success = False
        self.directory_path = ""
        self.total_files = 0
        self.success_count = 0
        self.fail_count = 0
        self.parsed_count = 0
        self.chunk_count = 0
        self.milvus_write_count = 0
        self.start_time: Optional[datetime] = None
        self.end_time: Optional[datetime] = None
        self.error_message = ""
        self.failed_files: Dict[str, str] = {}

    def increment_success_count(self):
        """增加成功计数"""
        self.success_count += 1

    def increment_fail_count(self):
        """增加失败计数"""
        self.fail_count += 1

    def add_failed_file(self, file_path: str, error: str):
        """添加失败文件"""
        self.failed_files[file_path] = error

    def get_duration_ms(self) -> int:
        """获取耗时（毫秒）"""
        if self.start_time and self.end_time:
            return int((self.end_time - self.start_time).total_seconds() * 1000)
        return 0

    def to_dict(self) -> Dict[str, Any]:
        """转换为字典"""
        return {
            "success": self.success,
            "directory_path": self.directory_path,
            "total_files": self.total_files,
            "success_count": self.success_count,
            "fail_count": self.fail_count,
            "parsed_count": self.parsed_count,
            "chunk_count": self.chunk_count,
            "milvus_write_count": self.milvus_write_count,
            "duration_ms": self.get_duration_ms(),
            "error_message": self.error_message,
            "failed_files": self.failed_files,
        }


class VectorIndexService:
    """向量索引服务 - 负责读取文件、生成向量、存储到 Milvus"""

    def __init__(self):
        """初始化向量索引服务"""
        self.upload_path = "./data/knowledge"
        self.supported_extensions = {".md", ".txt", ".pdf", ".docx"}
        logger.info("向量索引服务初始化完成")

    def index_directory(self, directory_path: Optional[str] = None) -> IndexingResult:
        """
        索引指定目录下的所有文件

        Args:
            directory_path: 目录路径（可选，默认使用配置的上传目录）

        Returns:
            IndexingResult: 索引结果
        """
        result = IndexingResult()
        result.start_time = datetime.now()

        try:
            # 使用指定目录或默认知识库目录
            target_path = directory_path if directory_path else self.upload_path
            dir_path = Path(target_path).resolve()

            if not dir_path.exists() or not dir_path.is_dir():
                raise ValueError(f"目录不存在或不是有效目录: {target_path}")

            result.directory_path = str(dir_path)

            # 递归获取所有支持的文件
            files = sorted(
                file_path
                for file_path in dir_path.rglob("*")
                if file_path.is_file() and file_path.suffix.lower() in self.supported_extensions
            )

            if not files:
                logger.warning(f"目录中没有找到支持的文件: {target_path}")
                result.total_files = 0
                result.success = True
                result.end_time = datetime.now()
                return result

            result.total_files = len(files)
            logger.info("扫描目录: {}", dir_path)
            logger.info("发现文件: {}个", len(files))

            # 遍历并索引每个文件
            seen_hashes: set[str] = set()
            for file_path in files:
                try:
                    file_hash = self._calculate_md5(file_path)
                    if file_hash in seen_hashes:
                        logger.info("跳过重复文件(MD5): {}", file_path)
                        continue
                    seen_hashes.add(file_hash)

                    chunk_count = self.index_single_file(str(file_path), file_hash=file_hash)
                    result.increment_success_count()
                    result.parsed_count += 1
                    result.chunk_count += chunk_count
                    result.milvus_write_count += chunk_count
                    logger.info(f"[OK] 文件索引成功: {file_path.name}")
                except Exception as e:
                    result.increment_fail_count()
                    result.add_failed_file(str(file_path), str(e))
                    logger.error(f"[FAIL] 文件索引失败: {file_path.name}, 错误: {e}")

            result.success = result.fail_count == 0
            result.end_time = datetime.now()
            try:
                milvus_manager.get_collection().flush()
                logger.info("Milvus collection 已 flush，实体数量统计已刷新")
            except Exception as e:
                logger.warning(f"Milvus flush 失败: {e}")

            logger.info(
                "索引完成: 扫描目录={}, 发现文件={}个, 成功解析={}个, 切分chunk={}个, 写入Milvus={}条, 失败={}个",
                dir_path,
                result.total_files,
                result.parsed_count,
                result.chunk_count,
                result.milvus_write_count,
                result.fail_count,
            )

            return result

        except Exception as e:
            logger.error(f"索引目录失败: {e}")
            result.success = False
            result.error_message = str(e)
            result.end_time = datetime.now()
            return result

    def index_single_file(self, file_path: str, file_hash: Optional[str] = None) -> int:
        """
        索引单个文件 (使用新的 LangChain 分割器)

        Args:
            file_path: 文件路径

        Raises:
            ValueError: 文件不存在时抛出
            RuntimeError: 索引失败时抛出
        """
        path = Path(file_path).resolve()

        if not path.exists() or not path.is_file():
            raise ValueError(f"文件不存在: {file_path}")

        logger.info(f"开始索引文件: {path}")

        try:
            # 1. 读取文件内容
            content = self._read_file_content(path)
            logger.info(f"读取文件: {path}, 内容长度: {len(content)} 字符")

            # 2. 删除该文件的旧数据（如果存在）
            normalized_path = path.as_posix()
            vector_store_manager.delete_by_source(normalized_path)

            # 3. 使用新的文档分割器
            documents = document_splitter_service.split_document(content, normalized_path)
            doc_hash = file_hash or self._calculate_md5(path)
            for doc in documents:
                doc.metadata["_file_md5"] = doc_hash
            logger.info(f"文档分割完成: {file_path} -> {len(documents)} 个分片")

            # 4. 添加文档到向量存储
            if documents:
                vector_store_manager.add_documents(documents)
                logger.info(f"文件索引完成: {file_path}, 共 {len(documents)} 个分片")
                return len(documents)
            else:
                logger.warning(f"文件内容为空或无法分割: {file_path}")
                return 0

        except Exception as e:
            logger.error(f"索引文件失败: {file_path}, 错误: {e}")
            raise RuntimeError(f"索引文件失败: {e}") from e

    def _read_file_content(self, path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix in {".md", ".txt"}:
            return path.read_text(encoding="utf-8")
        if suffix == ".pdf":
            return self._read_pdf_content(path)
        if suffix == ".docx":
            return self._read_docx_content(path)
        raise ValueError(f"不支持的文件格式: {suffix}")

    def _read_pdf_content(self, path: Path) -> str:
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise RuntimeError("缺少 PDF 解析依赖 pypdf，请先安装 pypdf") from exc

        reader = PdfReader(str(path))
        pages = []
        for index, page in enumerate(reader.pages, 1):
            text = page.extract_text() or ""
            if text.strip():
                pages.append(f"\n\n--- page {index} ---\n{text}")
        return "\n".join(pages)

    def _read_docx_content(self, path: Path) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("缺少 DOCX 解析依赖 python-docx，请先安装 python-docx") from exc

        document = Document(str(path))
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    def _calculate_md5(self, path: Path) -> str:
        md5 = hashlib.md5()
        with path.open("rb") as file:
            for chunk in iter(lambda: file.read(1024 * 1024), b""):
                md5.update(chunk)
        return md5.hexdigest()


# 全局单例
vector_index_service = VectorIndexService()
